#!/usr/bin/env python3
"""
Generate the downloadable, fill-in-the-blanks quant CV template.

Mirrors the house one-page UK format used across the site: Arial throughout,
a centred bold name, bold section headers in caps, each entry showing the
institution (bold, left) and location/date (right, via a right tab stop),
an italic sub-line for the role, and compact bullet points.

Run:  python3 scripts/gen_cv_template.py
Out:  public/quant-cv-template.docx   (committed to the repo, served at /)
"""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

FONT = "Arial"
RIGHT_TAB_MM = 159  # A4 content width with ~1in margins

doc = Document()

# --- page + base style -------------------------------------------------------
sec = doc.sections[0]
sec.page_height = Mm(297)
sec.page_width = Mm(210)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Mm(18))

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.space_before = Pt(0)


def _arial(run, size, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # force Arial for complex-script fallback too
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)


def _right_tab(p):
    p.paragraph_format.tab_stops.add_tab_stop(Mm(RIGHT_TAB_MM), WD_TAB_ALIGNMENT.RIGHT)


def name(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _arial(p.add_run(text), 20, bold=True)


def contact(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _arial(p.add_run(text), 10)


def section(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _arial(p.add_run(text.upper()), 11, bold=True)
    # thin underline rule beneath the header
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pbdr.append(bottom)
    pPr.append(pbdr)


def split_line(left, right, *, size, left_bold, left_italic, right_bold, right_italic):
    """One paragraph: left text, a right-aligned tab, then right text."""
    p = doc.add_paragraph()
    _right_tab(p)
    _arial(p.add_run(left), size, bold=left_bold, italic=left_italic)
    if right:
        p.add_run("\t")
        _arial(p.add_run(right), size, bold=right_bold, italic=right_italic)
    return p


def entry(institution, place_or_date, role=None, date=None, *, edu=False):
    # header line: institution (bold) ........ place/date (right)
    split_line(
        institution, place_or_date,
        size=11, left_bold=True, left_italic=False,
        right_bold=not edu, right_italic=edu,
    )
    # optional sub-line: role (italic) ........ date (italic)
    if role is not None:
        split_line(
            role, date or "",
            size=10, left_bold=False, left_italic=True,
            right_bold=False, right_italic=True,
        )


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        _arial(p.add_run(it), 10)


# --- content (placeholders only — replace everything in [brackets]) ----------
name("JOHN DOE")
contact("+44 7000 000000  |  1 Example Street, Anytown, AB1 2CD  |  JOHN.DOE@EMAIL.COM")

section("Education")
entry("Example University, MMath Mathematics", "Sep 20XX – Jul 20XX", edu=True)
bullets([
    "Predicted / achieved First Class Honours (XX% overall).",
    "Year [N] modules: [Probability] (XX%), [Stochastic Processes] (XX%), [Computational X] (XX%).",
    "Relevant coursework: [a project or dissertation in one line — what you built and the result].",
])
entry("Example Sixth Form College", "Sep 20XX – Aug 20XX", edu=True)
bullets([
    "A Levels: Mathematics (A*), Further Mathematics (A*), [Third subject] (A*).",
])
entry("Example Secondary School", "Sep 20XX – Aug 20XX", edu=True)
bullets([
    "GCSEs: [N] grade 9s including Mathematics, Sciences and Further Mathematics.",
])

section("Work Experience")
entry("Example Firm", "City, Country", "Insight / Internship Programme", "Month 20XX")
bullets([
    "[What you did, with a number: e.g. built a model that did X, reducing Y by Z%.]",
    "[A second concrete contribution — start with a verb, quantify the outcome.]",
    "[A result a reader can picture — what changed because you were there.]",
])
entry("Example Tutoring Company", "Remote, Country", "Part-time Maths Tutor", "Month 20XX – Present")
bullets([
    "Taught [N] GCSE/A-level students; [outcome — grades improved, retention, etc.].",
    "[Planned and delivered tailored lessons — one line on impact.]",
])
entry("Example Employer", "City, Country", "Part-time Role", "Month 20XX – Month 20XX")
bullets([
    "[An achievement, ideally with a superlative or a number.]",
    "[A responsibility that shows reliability or initiative.]",
])

section("Positions of Responsibility")
entry("Example Society / Sports Club", "City, Country", "Committee Role (e.g. President)", "Month 20XX – Present")
bullets([
    "[Elected/appointed how, and what you were responsible for.]",
    "[Something you organised or improved, with the result.]",
])
entry("Example School", "City, Country", "Leadership Role (e.g. Head Student)", "Month 20XX – Month 20XX")
bullets([
    "[A leadership or public-speaking responsibility, in one line.]",
])

section("Achievements, Interests, and Skills")
bullets([
    "Programming: Python (NumPy, pandas), [C++ / SQL / R], LaTeX, Git.",
    "Competitions: [UKMT Gold / Senior Kangaroo / olympiad], [other measurable result].",
    "Games: [chess rating / poker / bridge — anything with a concrete level].",
    "Languages: English (native), [other] ([level]).",
    "Interests: [music grade, sport to a level, anything that shows commitment and ranking].",
])

out = os.path.join(os.path.dirname(__file__), "..", "public", "quant-cv-template.docx")
doc.save(out)
print("wrote", os.path.normpath(out))
